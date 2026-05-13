"""生成 UniMind.ai 商业化启动方案 .docx 文档"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# 页边距
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return h

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_with_style(headers, rows, col_widths=None):
    """创建统一样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '2563EB')

    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F1F5F9')

    doc.add_paragraph()
    return table

# ================================================
# 封面
# ================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('UniMind.ai')
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(37, 99, 235)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('AI 驱动的 431 金融备考平台\n商业化启动方案')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(55, 65, 81)

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(f'版本：v1.0\n日期：{datetime.date.today().strftime("%Y年%m月%d日")}\n密级：内部资料')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(107, 114, 128)

doc.add_page_break()

# ================================================
# 目录页
# ================================================
add_heading_styled('目录', level=1)
toc_items = [
    '一、产品定位与核心价值主张',
    '二、目标用户画像',
    '三、三版定价体系与功能矩阵',
    '四、为什么这样定价——竞品对标与逻辑',
    '五、冷启动：怎么找到第一个付费客户',
    '六、小红书 / 抖音内容获客策略',
    '七、第一笔钱的 3 条最短路径',
    '八、90 天商业化行动路线图',
    '九、融资与退出路径',
    '十、常见问题预答',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(12)

doc.add_page_break()

# ================================================
# 一、产品定位
# ================================================
add_heading_styled('一、产品定位与核心价值主张', level=1)

add_heading_styled('1.1 一句话定位', level=2)
doc.add_paragraph(
    'UniMind.ai 是一个面向 431 金融考研培训的 AI 提效工具：'
    '用大模型代替人工出题，用算法替代盲目刷题，让独立教师一个人就具备教研组级别的输出能力。'
)

add_heading_styled('1.2 我们卖的不是"平台"，卖的是"产出"', level=2)
doc.add_paragraph('机构/教师购买的不是软件本身，而是：')
bullets = [
    '高质量题目（AI 三智能体对抗出题，质量接近真人命题组）',
    '个性化复习方案（Memorix 算法精准预测遗忘点，比 FSRS 更准）',
    '学生薄弱点可见性（学情看板替代 Excel，一眼看到谁在哪里卡住了）',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

add_heading_styled('1.3 核心差异化', level=2)
add_table_with_style(
    ['维度', '传统机构现状', 'UniMind'],
    [
        ['出题', '老师手动出题，1道题1-2小时', 'AI 10秒出题，三轮对抗保证质量'],
        ['复习安排', '统一布置，学霸学渣同一套', 'Memorix 为每个学生生成不同计划'],
        ['学情掌握', 'Excel 表格，月底才汇总', '实时仪表盘，今天谁没学会一目了然'],
        ['题库更新', '跟考纲跟不动，题目老化', 'AI 按新考点持续生成新题'],
    ],
)

doc.add_page_break()

# ================================================
# 二、目标用户画像
# ================================================
add_heading_styled('二、目标用户画像', level=1)

doc.add_paragraph(
    '我们将目标用户分为三个层级，对应三个版本。核心逻辑是：'
    '从独立教师切入（决策快、口碑效应强），再向机构渗透（客单价高、粘性强）。'
)

add_heading_styled('2.1 独立教师（Solo 版目标用户）', level=2)
profile_solo = [
    '画像：1-3 人的小团队或个人教师，专门辅导 431 金融考生',
    '痛点：出题耗时巨大，一个人要干教研组所有活；学生多了就顾不过来',
    '付费意愿：愿意为提效工具付费，但预算有限（¥200-500/月）',
    '决策特征：个人决策，不需要审批，试用觉得好当场就买',
    '市场规模：全国 431 考研辅导个人教师预估 3,000-5,000 人',
    '触达方式：小红书、知乎、考研微信群、431 备考社区',
]
for item in profile_solo:
    doc.add_paragraph(item, style='List Bullet')

add_heading_styled('2.2 小型辅导工作室 / 中小机构（Plus 版目标用户）', level=2)
profile_team = [
    '画像：5-20 人团队，有 100-300 学生，有 1-2 个线下教学点',
    '痛点：教学管理混乱，不同老师的题目质量参差不齐，缺乏数据化运营',
    '付费意愿：愿意为系统化工具付费（¥1,000-3,000/月），但要看到可量化的 ROI',
    '决策特征：需要创始人/校长拍板，决策周期 1-2 周',
    '触达方式：行业会议、考研培训机构圈子、转介绍',
]
for item in profile_team:
    doc.add_paragraph(item, style='List Bullet')

add_heading_styled('2.3 大型连锁考研机构（Pro 版目标用户）', level=2)
profile_ent = [
    '画像：跨城市连锁，500+ 学生，自有教研团队',
    '痛点：题库标准化、跨校区教学质量一致、品牌化输出',
    '付费意愿：预算充足（¥5,000-15,000/月），但需要定制 + SLA + 专属服务',
    '决策特征：招标式采购，需要商务谈判 + 演示 + 试用 + 合同，周期 1-3 月',
    '触达方式：商务 BD、行业展会、招投标信息',
]
for item in profile_ent:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ================================================
# 三、定价体系
# ================================================
add_heading_styled('三、三版定价体系与功能矩阵', level=1)

doc.add_paragraph(
    '核心定价原则：'
    '（1）独立教师版是"引流产品"，定价要低到让个人教师无需犹豫；'
    '（2）AI 出题能力全版本开放，因为它是核心价值，不应该被锁在最高版；'
    '（3）机构版溢价来自协作、管理、分析能力，而非基础功能；'
    '（4）学员数上限是天然的升级杠杆——教师做大了自然会升。'
)

add_heading_styled('3.1 定价总览', level=2)

add_table_with_style(
    ['', 'Solo 独立教师版', 'Team 机构版', 'Enterprise 企业版'],
    [
        ['月付价格', '¥299/月', '¥1,299/月', '¥3,999/月'],
        ['年付价格', '¥199/月（省 33%）', '¥999/月（省 23%）', '¥2,999/月（省 25%）'],
        ['学员上限', '30 人', '150 人', '不限'],
        ['教师账号', '1 个', '5 个', '不限'],
        ['免费试用', '14 天全功能', '14 天全功能', '定制 Demo'],
        ['核心定位', '个人提效工具', '团队协作平台', '企业级解决方案'],
    ],
)

add_heading_styled('3.2 完整功能矩阵', level=2)

add_table_with_style(
    ['功能模块', '具体功能', 'Solo', 'Team', 'Enterprise'],
    [
        ['题库管理', '手动创建/编辑/导入题目', '✅', '✅', '✅'],
        ['题库管理', '题型分类与标签系统', '✅', '✅', '✅'],
        ['AI 智能出题', '基于考点 AI 生成题目', '✅', '✅', '✅'],
        ['AI 智能出题', '三智能体对抗质量保障', '✅', '✅', '✅'],
        ['AI 智能出题', '批量出题（一次生成 N 道）', '✅', '✅', '✅'],
        ['组卷考试', '手动组卷 + 自动组卷', '✅', '✅', '✅'],
        ['组卷考试', '在线考试 + 自动批改', '✅', '✅', '✅'],
        ['基础统计', '学生成绩/正确率统计', '✅', '✅', '✅'],
        ['自适应复习', 'Memorix 个性化复习计划', '✅', '✅', '✅'],
        ['自适应复习', '遗忘曲线预测与干预', '✅', '✅', '✅'],
        ['课程管理', '视频课程上传与管理', '❌', '✅', '✅'],
        ['课程管理', 'AI 视频大纲时间戳打点', '❌', '✅', '✅'],
        ['知识图谱', '知识点可视化网络', '❌', '✅', '✅'],
        ['知识图谱', '学生个人掌握度图谱', '❌', '✅', '✅'],
        ['在线答疑', 'FAQ 系统 + AI 自动应答', '❌', '✅', '✅'],
        ['团队协作', '多教师账号与权限', '❌', '✅', '✅'],
        ['团队协作', '题目审核工作流', '❌', '✅', '✅'],
        ['学情看板', '班级/个人实时数据看板', '❌', '❌', '✅'],
        ['学情看板', 'AI 学习建议与预警', '❌', '❌', '✅'],
        ['高级分析', '教师效率分析', '❌', '❌', '✅'],
        ['高级分析', '跨班级横向对比', '❌', '❌', '✅'],
        ['学习房间', '实时在线学习室 WebSocket', '❌', '❌', '✅'],
        ['品牌定制', '机构自定义 Logo/域名', '❌', '❌', '✅'],
        ['品牌定制', '白标部署（独立域名）', '❌', '❌', '✅'],
        ['API 对接', 'REST API 接口开放', '❌', '❌', '✅'],
        ['专属服务', '专属客户成功经理', '❌', '❌', '✅'],
        ['专属服务', '定制功能开发协商', '❌', '❌', '✅'],
    ],
)

add_heading_styled('3.3 为什么 Solo 版包含 AI 出题和 Memorix', level=2)
doc.add_paragraph(
    '这是核心差异化决策。理由是：'
)
reasons = [
    '如果 Solo 版没有 AI 出题，独立教师进来看到的是一个"空壳题库系统"——没有内容，没有价值感知，秒退。',
    '独立教师是口碑的起点。他用了觉得好，会在他所在的机构、教研圈子里传播。如果他把最好的功能锁在机构版，他永远不会推荐。',
    'Memorix 的壁垒不在算法本身，而在于"使用越多数据越准"。Solo 用户的答题数据是我们模型进化的燃料。',
    '独立教师做到 30 个学生封顶后，自然会升级到 Team 版——这是产品驱动的增长，不需要销售推动。',
]
for r in reasons:
    doc.add_paragraph(r, style='List Number')

add_heading_styled('3.4 超出学员上限的处理', level=2)
doc.add_paragraph(
    'Solo 版超过 30 人 → 系统提示升级 Team，但给予 7 天宽限期。'
    'Team 版超过 150 人 → 系统提示升级 Enterprise，或按 ¥8/人/月加购名额。'
    '这确保了不会因为"突然多招了几个学生"就断掉服务，同时制造自然的升级节点。'
)

doc.add_page_break()

# ================================================
# 四、定价逻辑
# ================================================
add_heading_styled('四、为什么这样定价——竞品对标与逻辑', level=1)

add_heading_styled('4.1 用户视角的 ROI 计算', level=2)

doc.add_paragraph('独立教师视角：')
roi_solo = [
    '一个独立教师每月收入：20 学生 × ¥400/人 = ¥8,000',
    '每月出题时间：每天 2 小时 × 22 天 = 44 小时',
    'UniMind Solo：¥299/月，省下 30+ 小时出题时间',
    '时薪对比：¥299 ÷ 30h = ¥10/小时 → 比请兼职还便宜 10 倍',
]
for r in roi_solo:
    doc.add_paragraph(r, style='List Bullet')

doc.add_paragraph('小型机构视角：')
roi_team = [
    '一个机构每月收入：150 学生 × ¥500/人 = ¥75,000',
    '传统教研成本：至少 1 个全职教研老师 → ¥6,000-10,000/月',
    'UniMind Team：¥1,299/月，相当于一个教研老师工资的 1/5',
    '还不算题目质量提升带来的续费率提升',
]
for r in roi_team:
    doc.add_paragraph(r, style='List Bullet')

add_heading_styled('4.2 竞品价格对标', level=2)
add_table_with_style(
    ['竞品', '类型', '起步价', '我们相比'],
    [
        ['问卷星/考试云', '通用考试工具', '¥199-599/月', '我们有 AI 出题，它们只有组卷'],
        ['小鹅通', '知识付费 SaaS', '¥4,999-19,999/年', '我们是垂直 431 的专用工具'],
        ['学科网', 'K12 题库平台', '按校定价 ¥5,000+', '我们聚焦考研，更垂直'],
        ['自研系统', '大机构自建', '¥50,000-200,000 一次性', '我们 1/10 价格，持续更新'],
    ],
)

doc.add_paragraph(
    '结论：Solo ¥299 低于通用 SaaS 的中间档，Team ¥1,299 低于小鹅通入门版，Enterprise ¥3,999/月低于雇一个初级工程师。'
    '每个价位都有明确的"便宜感"。'
)

doc.add_page_break()

# ═══════════════════════════════════════
# 五、冷启动（纯社交媒体版）
# ═══════════════════════════════════════
add_heading_styled('五、冷启动：用社交媒体找到第一个付费客户', level=1)

doc.add_paragraph(
    '事实：你没有品牌、没有案例、没有用户评价。但你有社交媒体——'
    '一个 0 粉丝的新号，发第一篇就可能被推荐给 5,000 个精准用户。'
    '冷启动的核心不是"找客户"，是"让客户找到你"。'
)

add_heading_styled('5.1 为什么社交媒体比私信更好', level=2)
cold_reasons = [
    '一篇笔记发出去，持续获客 3-6 个月（长尾搜索流量），私信一次只能触达一个人',
    '用户主动私信你的，已经通过内容建立了信任，转化率远高于你主动私信陌生人',
    '你不需要任何"人脉"——小红书的算法不看你认识谁，看你的内容好不好',
    '社交媒体的反馈是数量化的：阅读量、收藏量、私信量，你可以快速判断什么内容有效',
]
for r in cold_reasons:
    doc.add_paragraph(r, style='List Bullet')

add_heading_styled('5.2 第一步：账号就位', level=2)
doc.add_paragraph('在小红书注册一个账号，做好这三件事：')
cold_step1 = [
    '账号名：建议「AI 提效的 431 老师」或「431 出题机器人」——直击身份 + 关键词，方便搜索命中',
    '头像：用 AI 生成一个专业但不严肃的头像（见宣传图 Prompt 文档）',
    '简介："教 431 金融的第 X 年 | 用 AI 把出题效率提升 10 倍 | 正在做一个让老师不用熬夜出题的工具 | 私信「体验」免费试用"',
]
for s in cold_step1:
    doc.add_paragraph(s, style='List Bullet')

add_heading_styled('5.3 第二步：发布前 4 篇内容', level=2)
doc.add_paragraph('不要发产品广告。前 4 篇的目标是建立"这个人懂 431"的专业认知：')

add_table_with_style(
    ['篇序', '选题', '类型', '为什么选这个'],
    [
        ['第 1 篇', '"分析 10 年 431 真题，发现这 5 个考点几乎每年都考"', '数据揭秘', '最容易拿收藏和搜索流量'],
        ['第 2 篇', '"带 30 个学生，我每天出题到凌晨 1 点"', '痛点共鸣', '让目标用户产生"这说的不就是我吗"'],
        ['第 3 篇', '"同一道题：手工出题 2h vs AI 出题 10s"', '工具展示', '对比冲击，展示你的产品解决什么问题'],
        ['第 4 篇', '"为什么刷 3000 题不如别人刷 300 题"', '认知输出', '建立方法论层面的权威'],
    ],
)

doc.add_paragraph('每篇的结尾都加一句：')
doc.add_paragraph(
    '"我最近在做 AI 出题工具，把这些方法做成了自动化。想试试可以私信我。"'
)
doc.add_paragraph('不要硬推，不要放链接（小红书会限流），让用户主动来找你。')

add_heading_styled('5.4 第三步：等私信，快速回复', level=2)
doc.add_paragraph(
    '发完 4 篇后，每天早中晚各检查一次私信。当有人私信"想了解"时，回复模板：'
)
doc.add_paragraph(
    '"你好！我做了个 AI 出题工具，专门给 431 辅导老师用的。核心功能是输入考点→AI 自动生成高质量题目，'
    '还有自适应复习算法帮学生记住做过的题。目前有 14 天免费试用，不用绑卡。'
    '注册链接是 unimind.ai/register，用完之后告诉我感受～"'
)

add_heading_styled('5.5 第四步：深度 Onboarding + 收钱', level=2)
doc.add_paragraph('有人试用之后：')
onboarding = [
    '主动加微信或开视频，30 分钟共享屏幕演示完整流程：出题→组卷→学生做题→看数据',
    '帮他导入他现有的题目（哪怕手动帮他录 20 道）',
    '第一周每隔 2 天问一次"用的怎么样，有什么不方便的"',
    '第 13 天发消息："你的 14 天试用明天到期了～觉得有用的话，现在早期用户 ¥199/月（年付）锁定终身不涨价，要不要试试？"',
]
for o in onboarding:
    doc.add_paragraph(o, style='List Bullet')

doc.add_paragraph(
    '目标：第一个月发 12 篇内容 → 获得 3-5 个私信咨询 → 2-3 个试用 → 1 个付费。这就够了。'
    '详见《UniMind 社交媒体获客方案》文档获取完整内容日历和 4 周执行计划。'
)

doc.add_page_break()

# ================================================
# 六、小红书 / 抖音策略
# ================================================
add_heading_styled('六、小红书 / 知乎 / 抖音获客策略', level=1)

doc.add_paragraph(
    '核心原则：不要卖产品，卖"认知"。'
    'B 端客户不会因为看到一个广告就下单，但会因为持续看到你的专业输出而建立信任。'
    '当他在某个深夜为出题头疼时，会想起"那个做 AI 出题的人"，然后主动来找你。'
)

doc.add_paragraph(
    '本节为概览。完整策略（含 4 周内容日历、私信回复 SOP、转化漏斗指标、FAQ）'
    '见单独文档：《UniMind 社交媒体获客方案》。'
)

add_heading_styled('6.1 主战场：小红书', level=2)
doc.add_paragraph('五种内容类型循环发布：')
add_table_with_style(
    ['内容类型', '示例标题', '频率', '目的'],
    [
        ['数据揭秘', '"10年431真题，这5个考点考了87次"', '每周 1 篇', '搜索流量 + 专业信任'],
        ['痛点共鸣', '"带30个学生，每天出题到凌晨1点"', '每周 1 篇', '吸引目标用户'],
        ['工具展示', '"同一道题，手工2h vs AI 10s"', '每周 1 篇', '产品种草'],
        ['认知输出', '"为什么刷3000题不如刷300题"', '每 2 周 1 篇', '思想领袖定位'],
        ['社会证明', '"第一个用户用了2周后的真实反馈"', '有案例就发', '降低决策风险'],
    ],
)

add_heading_styled('6.2 辅助渠道：知乎 + 抖音', level=2)
doc.add_paragraph(
    '知乎：回答"431 金融学综合怎么准备"等高频搜索问题，一篇回答持续引流 2 年以上。'
    '抖音/B 站：用"手工出题 vs AI 出题"的对比视频制造冲击感，15-30 秒横版 + 竖版同时发。'
)

doc.add_page_break()

# ================================================
# 七、第一笔钱的三条路径
# ================================================
add_heading_styled('七、第一笔钱的 3 条最短路径', level=1)

add_heading_styled('路径 A：卖题库服务（最快，1-2 周见效）', level=2)
path_a = [
    '做法：找到 1 个 431 辅导老师，用 AI 为 TA 生成一套完整的章节题库（比如"货币银行学 200 题"），按套收费 ¥500-1,500',
    '优点：不需要对方用你的平台，交付一个 Word/PDF 就行，决策成本极低',
    '缺点：一锤子买卖，没有复购',
    '适合：快速验证 AI 出题质量 + 建立信任关系',
]
for p in path_a:
    doc.add_paragraph(p, style='List Bullet')

add_heading_styled('路径 B：独立教师 SaaS 订阅（推荐，2-4 周见效）', level=2)
path_b = [
    '做法：找到 3-5 个独立教师，14 天免费试用 → 深度 Onboarding → 转化 Solo 年付 ¥199/月',
    '优点：MRR（月度经常性收入），持续复购，用户反馈驱动产品迭代',
    '缺点：需要比路径 A 更多的客户成功投入',
    '目标：3 个付费用户 = 月收入约 ¥600，更重要的是有了 3 个活案例',
]
for p in path_b:
    doc.add_paragraph(p, style='List Bullet')

add_heading_styled('路径 C：白标定制（客单价最高，1-2 月见效）', level=2)
path_c = [
    '做法：找一家中小机构，帮他们把整套线上备考系统做起来，打他们的 Logo，用他们的域名，收费 ¥20,000-50,000',
    '优点：一单收入可观，深度了解机构需求',
    '缺点：定制化消耗大，可能影响产品化方向',
    '适合：同时推进，但不作为主方向',
]
for p in path_c:
    doc.add_paragraph(p, style='List Bullet')

add_heading_styled('建议优先级', level=2)
doc.add_paragraph('B（SaaS 订阅）> A（卖题库）> C（白标定制）。理由：B 是可持续的商业模式，A 是 B 的获客手段，C 是机会主义收入。')

doc.add_page_break()

# ================================================
# 八、90 天行动计划
# ================================================
add_heading_styled('八、90 天商业化行动路线图', level=1)

add_heading_styled('第 1-2 周：产品就绪', level=2)
week12 = [
    '完成版本权限控制系统开发（参考架构文档 BUSINESS_VERSION_PERMISSION_ARCHITECTURE.md）',
    '完成三版定价页面前端（Solo / Team / Enterprise）',
    '用 AI 生成 1 套 431 完整题库（至少 500 题），确保 Demo 时有内容可展示',
    '用宣传图 Prompt 生成 5 张小红书封面图',
    '完成产品演示视频第一版（录屏 + AI 配音，参考视频 Prompt 文档）',
]
for item in week12:
    doc.add_paragraph(item, style='List Bullet')

add_heading_styled('第 3-4 周：获客启动（纯社交媒体）', level=2)
week34 = [
    '小红书注册账号，设置好名称、头像、简介（按 5.2 步骤执行）',
    '发布前 4 篇内容：数据揭秘 + 痛点共鸣 + 工具展示 + 认知输出（按前 4 周内容日历执行）',
    '每天早中晚检查私信，有咨询 5 分钟内回复',
    '为每位试用用户做深度 Onboarding（30 分钟视频通话）',
    '收集反馈，快速迭代（本周发现的 Bug 本周修）',
    '同步开始知乎答题：每周 1-2 个高质量回答',
    '目标：3-5 个私信咨询 → 2-3 个试用 → 争取第 1 个付费用户',
]
for item in week34:
    doc.add_paragraph(item, style='List Bullet')

add_heading_styled('第 5-8 周：验证 PMF', level=2)
week58 = [
    '持续小红书内容输出（每周 2-3 篇）+ 知乎答题（每周 1-2 个）',
    '根据前 4 周数据，判断哪种内容类型效果最好，加大该类型产出',
    '前 1-2 个付费用户如果满意，请他们在小红书/朋友圈发使用体验',
    '开始接触小型机构（3-5 家），通过内容吸引来咨询的优先跟进',
    '收集至少 10 个用户反馈，迭代产品 2 个版本',
    '目标：3-5 个付费 Solo 用户 + 1 个 Team 用户',
]
for item in week58:
    doc.add_paragraph(item, style='List Bullet')

add_heading_styled('第 9-12 周：增长奠基', level=2)
week912 = [
    '跑通小红书内容→私信→试用→付费的完整转化链路，记录各环节转化率',
    '根据前 2 个月数据，决定：继续自造血 or 接触投资人 or 接触潜在收购方',
    '如果自造血：将内容创作 SOP 化（模板 + AI 辅助），把每周 6h 压缩到 3h',
    '抖音/B 站开始发第一条视频（复用小红书素材 + 录屏即可）',
    '如果融资：整理数据包（MRR、用户数、留存率、获客成本），写 BP',
    '目标：MRR ¥3,000+ 或拿到 TS（投资意向书）',
]
for item in week912:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ================================================
# 九、融资与退出
# ================================================
add_heading_styled('九、融资与退出路径', level=1)

add_heading_styled('9.1 不同阶段的估值逻辑', level=2)

add_table_with_style(
    ['阶段', '条件', '典型估值', '适合的资本类型'],
    [
        ['Pre-Seed', '有产品 + 1-2 个付费客户', '¥200-500 万', '天使投资人 / 考研行业个人'],
        ['Seed', 'MRR ¥2 万+，月增长 20%+', '¥1,000-3,000 万', '教育科技 VC / 加速器'],
        ['自造血', 'MRR ¥5 万+，盈利', '不融资，拿分红', '自己'],
    ],
)

add_heading_styled('9.2 卖给大机构的退出逻辑', level=2)
doc.add_paragraph(
    '考研培训行业有一个特点：大量中小机构缺乏技术能力，而头部机构（如文都、跨考、海文等）有收购意愿。'
)
exit_points = [
    '最佳退出时机：有 20+ 付费机构 + 月活学生 2,000+ → 战略价值明确',
    '潜在收购方：大型考研连锁机构（需要技术能力补全）、在线教育平台（需要垂直内容）、金融培训公司（需要 431 产品线）',
    '退出估值：通常按年营收的 3-5 倍 或 ARR（年化经常性收入）的 5-10 倍',
    '如果 MRR ¥5 万 → ARR ¥60 万 → 估值 ¥300-600 万',
]
for e in exit_points:
    doc.add_paragraph(e, style='List Bullet')

add_heading_styled('9.3 建议路径', level=2)
doc.add_paragraph(
    '先做自造血 6 个月。如果 MRR 过 ¥2 万且持续增长 → 决定是继续自己做（利润驱动）还是融资（规模驱动）。'
    '如果做了 6 个月 MRR 仍低于 ¥5,000 → 考虑两条路：（1）作为 Side Project 维持，不 All-in；'
    '（2）打包卖给行业内的技术需求方。'
)

doc.add_page_break()

# ================================================
# 十、常见问题预答
# ================================================
add_heading_styled('十、常见问题预答', level=1)

faqs = [
    (
        'Q: Solo 版 ¥299/月太便宜了，能赚钱吗？',
        'A: Solo 版的目的不是赚钱，是获客 + 口碑。一个 Solo 教师如果做大了，自然会升级到 Team。'
        '同时，Solo 用户贡献的答题数据在训练 Memorix 模型，这是长期壁垒。'
        '真正赚钱的是 Team（¥1,299/月）和 Enterprise（¥3,999/月）。'
    ),
    (
        'Q: AI 出题在 Solo 版开放，不怕机构版卖不动吗？',
        'A: 机构的购买决策不是因为"AI 出题"这一个功能，而是因为协作、管理、分析、品牌定制这些 Solo 版给不了的。'
        '一个开火锅店的人不会因为家用电磁炉好用就不买商用厨房设备。'
    ),
    (
        'Q: 如果竞争对手抄怎么办？',
        'A: 三件事无法被快速复制：'
        '（1）Memorix 算法需要大量真实答题数据训练，数据越多越准；'
        '（2）三智能体对抗出题的 Prompt 工程是经过大量调优的；'
        '（3）对 431 考纲的理解深度——这不是通用 AI 能替代的。'
        '所以不要怕竞争，要怕的是不动。先跑出 100 个付费用户，你的数据壁垒就建起来了。'
    ),
    (
        'Q: 免费试用 14 天后用户不付费怎么办？',
        'A: 如果大部分用户试用后不付费，不是定价问题，是产品没解决真痛点。'
        '复盘方向：（1）他试用期间用了几次 AI 出题？（2）他有没有把生成的题目真的发给学生做？（3）学生做题数据有没有进系统？'
        '如果回答都是"没有"，说明 Onboarding 没做到位，他没真正用起来。'
    ),
    (
        'Q: 我应该辞职 All-in 吗？',
        'A: 不要。在当前阶段（0 付费 → 第 1 个付费），UniMind 应该是一个 Side Project。'
        '等到月收入能覆盖你基本生活开销（比如 ¥8,000-10,000/月）的那一天，再考虑全职。'
        '这条路上死掉的人，99% 是在产品还没 PMF 的时候就 All-in 了。'
    ),
]

for q, a in faqs:
    add_heading_styled(q, level=2)
    doc.add_paragraph(a)

# ── 尾页 ──
doc.add_page_break()
for _ in range(10):
    doc.add_paragraph()

end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = end.add_run('— 文档结束 —')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(156, 163, 175)

end2 = doc.add_paragraph()
end2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = end2.add_run('UniMind.ai 商业化启动方案 v1.0')
run2.font.size = Pt(10)
run2.font.color.rgb = RGBColor(156, 163, 175)

# ── 保存 ──
output_path = '/Users/eular/Desktop/官网0215/docs/UniMind商业化启动方案.docx'
doc.save(output_path)
print(f'✅ 文档已生成：{output_path}')
